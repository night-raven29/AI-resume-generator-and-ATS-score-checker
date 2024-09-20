import streamlit as st
import requests
from io import BytesIO
from docx import Document
import pdfplumber
import re

# Function to extract text from a DOCX file
def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    return '\n'.join([paragraph.text for paragraph in doc.paragraphs])

# Function to extract text from a PDF file
def extract_text_from_pdf(pdf_file):
    with pdfplumber.open(pdf_file) as pdf:
        text = ''.join([page.extract_text() for page in pdf.pages])
    return text

# Function to extract key details from text using simple regex (or extend using NLP if needed)
def extract_details_from_text(resume_text):
    name = re.findall(r"Name[:\-\s]*([A-Za-z\s]+)", resume_text)
    email = re.findall(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", resume_text)
    phone = re.findall(r"\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}", resume_text)

    # Return extracted details
    return {
        "name": name[0] if name else "Name not found",
        "email": email[0] if email else "Email not found",
        "phone": phone[0] if phone else "Phone number not found"
    }

# Function to create a .docx resume based on the fields
def create_resume_docx(details):
    doc = Document()
    doc.add_heading(details['name'], 0)
    doc.add_paragraph(details['job_title'])

    doc.add_heading('Personal Details', level=1)
    doc.add_paragraph(f"Email: {details['email']}")
    doc.add_paragraph(f"Mobile: {details['mobile']}")
    doc.add_paragraph(f"LinkedIn: {details['linkedin']}")
    doc.add_paragraph(f"GitHub: {details['github']}")

    doc.add_heading('Education', level=1)
    doc.add_paragraph(details['education'])

    doc.add_heading('Work Experience', level=1)
    doc.add_paragraph(details['work_experience'])

    doc.add_heading('Skills', level=1)
    doc.add_paragraph(', '.join(details['skills']))

    doc.add_heading('Personal Projects', level=1)
    doc.add_paragraph(details['personal_projects'])

    doc.add_heading('Certificates', level=1)
    doc.add_paragraph(details['certificates'])

    doc.add_heading('Languages', level=1)
    doc.add_paragraph(details['languages'])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Function to send resume to ATS scoring API (API Deck) and get results
def get_ats_score_deck(file, api_key):
    url = "https://api.apideck.com/v1/ats/resume-score"  # Example endpoint, modify based on actual API
    headers = {
        'Authorization': f'Bearer {api_key}',
        'Content-Type': 'application/json'
    }
    
    files = {'resume': file}

    response = requests.post(url, headers=headers, files=files)
    
    if response.status_code == 200:
        return response.json()  # Assuming the API returns a JSON with ATS score and feedback
    else:
        return None

# Streamlit App
def resume_and_ats_checker():
    st.title("AI Resume Generator & ATS Score Checker (API Deck)")

    # Section 1: Resume Generation with Drag-and-Drop File Upload
    st.header("Resume Generator")

    # Drag-and-drop to upload existing resume
    uploaded_file = st.file_uploader("Drag and drop your existing resume to auto-fill (optional)", type=["docx", "pdf", "txt"])

    if uploaded_file is not None:
        if uploaded_file.type == "application/pdf":
            resume_text = extract_text_from_pdf(uploaded_file)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            resume_text = extract_text_from_docx(uploaded_file)
        else:
            resume_text = uploaded_file.read().decode('utf-8')

        extracted_details = extract_details_from_text(resume_text)
        st.success("Resume details auto-filled from uploaded file.")
    else:
        extracted_details = {}

    # Allow manual entry or overwrite of details
    name = st.text_input("Name", extracted_details.get('name', ''))
    job_title = st.text_input("Job Title")
    email = st.text_input("Email", extracted_details.get('email', ''))
    mobile = st.text_input("Mobile Number", extracted_details.get('phone', ''))
    linkedin = st.text_input("LinkedIn")
    github = st.text_input("GitHub")
    education = st.text_area("Education")
    work_experience = st.text_area("Work Experience")
    skills = st.text_input("Skills (comma separated)")
    personal_projects = st.text_area("Personal Projects")
    certificates = st.text_area("Certificates")
    languages = st.text_input("Languages (comma separated)")

    # Collect details
    details = {
        "name": name,
        "job_title": job_title,
        "email": email,
        "mobile": mobile,
        "linkedin": linkedin,
        "github": github,
        "education": education,
        "work_experience": work_experience,
        "skills": skills.split(',') if skills else [],
        "personal_projects": personal_projects,
        "certificates": certificates,
        "languages": languages
    }

    if st.button("Generate Resume"):
        # Generate resume .docx
        docx_buffer = create_resume_docx(details)
        st.download_button(
            label="Download Resume",
            data=docx_buffer,
            file_name=f"{name}_resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # Section 2: ATS Score Checker using API Deck
    st.header("ATS Score Checker with API Deck")
    uploaded_resume = st.file_uploader("Upload Generated Resume (.docx or .pdf)", type=["docx", "pdf"])

    if uploaded_resume:
        if st.button("Check ATS Score"):
            api_key = "your_apideck_api_key"  # Replace with your actual API Deck key
            ats_result = get_ats_score_deck(uploaded_resume, api_key)
            
            if ats_result:
                st.write(f"ATS Score: {ats_result['score']}%")
                st.write(f"Matched Keywords: {', '.join(ats_result['matched_keywords'])}")
                st.write(f"Suggestions: {ats_result['suggestions']}")
            else:
                st.error("Failed to get ATS score. Please try again.")

if __name__ == "__main__":
    resume_and_ats_checker()
